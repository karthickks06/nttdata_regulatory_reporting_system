"""Document parser sub-agent for extracting text from various document formats"""

from typing import Dict, Any, Optional, List
from pathlib import Path
import asyncio
from datetime import datetime

from app.tools.pdf_parser import parse_pdf
from app.tools.xml_parser import parse_xml
from app.core.config import settings


class DocumentParserAgent:
    """
    Sub-agent specialized in parsing documents (PDF, Word, Excel, XML).

    Capabilities:
    - Parse PDF documents (text, tables, metadata)
    - Parse Word documents (.docx)
    - Parse Excel spreadsheets (.xlsx)
    - Parse XML files
    - Extract structured data from regulatory documents
    """

    def __init__(self):
        self.name = "DocumentParser"
        self.supported_formats = [".pdf", ".docx", ".xlsx", ".xml", ".txt"]

    async def parse_document(
        self,
        file_path: str,
        document_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse document and extract content.

        Args:
            file_path: Path to the document file
            document_type: Optional document type hint

        Returns:
            Dictionary containing parsed content, metadata, and structure
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_extension = path.suffix.lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_extension}")

        try:
            if file_extension == ".pdf":
                result = await self._parse_pdf(path)
            elif file_extension == ".docx":
                result = await self._parse_docx(path)
            elif file_extension == ".xlsx":
                result = await self._parse_xlsx(path)
            elif file_extension == ".xml":
                result = await self._parse_xml(path)
            elif file_extension == ".txt":
                result = await self._parse_text(path)
            else:
                raise ValueError(f"Format not implemented: {file_extension}")

            # Add metadata
            result["metadata"] = {
                "file_path": str(path),
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "file_type": file_extension,
                "parsed_at": datetime.utcnow().isoformat(),
                "document_type": document_type
            }

            return result

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_path": str(path),
                "parsed_at": datetime.utcnow().isoformat()
            }

    async def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """Parse PDF document"""
        # Use PDF parser tool
        content = await parse_pdf(str(path))
        return {
            "success": True,
            "content": content.get("text", ""),
            "pages": content.get("pages", []),
            "tables": content.get("tables", []),
            "images": content.get("images", []),
            "metadata": content.get("metadata", {})
        }

    async def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """Parse Word document"""
        from docx import Document

        doc = Document(str(path))

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # Extract metadata
        core_properties = doc.core_properties
        metadata = {
            "author": core_properties.author,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "created": str(core_properties.created) if core_properties.created else None,
            "modified": str(core_properties.modified) if core_properties.modified else None
        }

        return {
            "success": True,
            "content": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": metadata
        }

    async def _parse_xlsx(self, path: Path) -> Dict[str, Any]:
        """Parse Excel spreadsheet"""
        from openpyxl import load_workbook

        workbook = load_workbook(str(path), read_only=True)

        sheets_data = {}
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            # Extract rows
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append(list(row))

            sheets_data[sheet_name] = rows

        workbook.close()

        return {
            "success": True,
            "sheets": sheets_data,
            "sheet_names": list(sheets_data.keys()),
            "content": str(sheets_data)
        }

    async def _parse_xml(self, path: Path) -> Dict[str, Any]:
        """Parse XML file"""
        result = await parse_xml(str(path))
        return {
            "success": True,
            "content": result.get("text", ""),
            "structure": result.get("structure", {}),
            "elements": result.get("elements", [])
        }

    async def _parse_text(self, path: Path) -> Dict[str, Any]:
        """Parse plain text file"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            "success": True,
            "content": content,
            "lines": content.split('\n')
        }

    async def extract_requirements(
        self,
        parsed_content: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Extract regulatory requirements from parsed document.

        Args:
            parsed_content: Parsed document content

        Returns:
            List of extracted requirements with metadata
        """
        requirements = []
        content = parsed_content.get("content", "")

        # Split into sections
        sections = content.split('\n\n')

        for idx, section in enumerate(sections):
            if len(section.strip()) < 50:
                continue

            # Check if section contains requirement keywords
            keywords = ["must", "shall", "should", "required", "mandatory", "obligation"]
            if any(keyword in section.lower() for keyword in keywords):
                requirements.append({
                    "requirement_text": section.strip(),
                    "section_number": idx + 1,
                    "is_mandatory": "must" in section.lower() or "shall" in section.lower(),
                    "extracted_at": datetime.utcnow().isoformat()
                })

        return requirements

    async def extract_entities(
        self,
        parsed_content: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Extract named entities from parsed document.

        Args:
            parsed_content: Parsed document content

        Returns:
            List of extracted entities
        """
        entities = []
        content = parsed_content.get("content", "")

        # Simple entity extraction (can be enhanced with NLP)
        # Extract regulatory references (e.g., "FCA PS21/23")
        import re

        # Pattern for regulatory references
        reg_pattern = r'\b(FCA|PRA|BOE|EBA)\s+[A-Z]{2}\d{2}/\d{2}\b'
        matches = re.findall(reg_pattern, content)

        for match in matches:
            entities.append({
                "entity_type": "REGULATORY_REFERENCE",
                "entity_value": match,
                "extracted_at": datetime.utcnow().isoformat()
            })

        return entities
